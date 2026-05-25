"""Word (.docx) document generator with Chinese support."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
from config import OUTPUT_DIR


def _set_chinese_font(run, font_name="微软雅黑", size=11):
    run.font.name = font_name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def create_doc(
    filename: str,
    title: str,
    sections: list[dict],
    subtitle: str = "",
) -> str:
    """Create a Word document.

    sections: list of dicts with keys:
      - heading (str, optional)
      - content (str) — body text, supports \\n for line breaks
      - style (str, optional) — 'normal'|'quote'|'blue' (default: normal)
    """
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "微软雅黑"
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(title)
    _set_chinese_font(run, size=18)
    run.bold = True

    if subtitle:
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_sub.add_run(subtitle)
        _set_chinese_font(run, size=11)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph("")  # spacer

    for sec in sections:
        if sec.get("heading"):
            h = doc.add_heading(level=2)
            run = h.add_run(sec["heading"])
            _set_chinese_font(run, size=14)
            run.bold = True

        text = sec.get("content", "")
        text_style = sec.get("style", "normal")

        for para_text in text.split("\n"):
            if not para_text.strip():
                doc.add_paragraph("")
                continue
            p = doc.add_paragraph()
            run = p.add_run(para_text)
            _set_chinese_font(run, size=11)
            if text_style == "blue":
                run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
            elif text_style == "quote":
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                run.italic = True

    path = os.path.join(OUTPUT_DIR, filename)
    doc.save(path)
    return path


def create_resume_doc(filename: str, content: dict) -> str:
    """Create a professional resume document.

    content keys: name, phone, email, target, summary, sections (list of
    {heading, items: [str]})
    """
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "微软雅黑"
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(content.get("name", ""))
    _set_chinese_font(run, size=20)
    run.bold = True

    # Contact line
    contact_parts = []
    if content.get("phone"):
        contact_parts.append(content["phone"])
    if content.get("email"):
        contact_parts.append(content["email"])
    if contact_parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(" | ".join(contact_parts))
        _set_chinese_font(run, size=10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    if content.get("target"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"目标岗位：{content['target']}")
        _set_chinese_font(run, size=10)
        run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)

    # Divider
    doc.add_paragraph("─" * 50)

    # Summary
    if content.get("summary"):
        h = doc.add_heading(level=2)
        run = h.add_run("个人简介")
        _set_chinese_font(run, size=13)
        run.bold = True
        p = doc.add_paragraph()
        run = p.add_run(content["summary"])
        _set_chinese_font(run, size=11)

    # Sections
    for sec in content.get("sections", []):
        h = doc.add_heading(level=2)
        run = h.add_run(sec["heading"])
        _set_chinese_font(run, size=13)
        run.bold = True
        for item in sec.get("items", []):
            p = doc.add_paragraph()
            run = p.add_run(item)
            _set_chinese_font(run, size=11)

    # Suggestions
    if content.get("suggestions"):
        doc.add_paragraph("")
        h = doc.add_heading(level=2)
        run = h.add_run("优化建议")
        _set_chinese_font(run, size=13)
        run.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
        for i, sug in enumerate(content["suggestions"], 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {sug}")
            _set_chinese_font(run, size=10)
            run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)

    path = os.path.join(OUTPUT_DIR, filename)
    doc.save(path)
    return path


def create_comparison_doc(
    filename: str,
    title: str,
    original_lines: list[str],
    revised_lines: list[str],
    reasons: list[str],
    overall_score: str = "",
    tips: list[str] | None = None,
) -> str:
    """Create a document with original vs revised comparison table."""
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "微软雅黑"
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    _set_chinese_font(run, size=18)
    run.bold = True

    if overall_score:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"整体评分：{overall_score}")
        _set_chinese_font(run, size=12)
        run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)

    # Polished full text
    h = doc.add_heading(level=2)
    run = h.add_run("润色终稿")
    _set_chinese_font(run, size=14)
    run.bold = True
    for line in revised_lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        _set_chinese_font(run, size=11)

    doc.add_paragraph("")

    # Comparison table
    h = doc.add_heading(level=2)
    run = h.add_run("逐句修改对照表")
    _set_chinese_font(run, size=14)
    run.bold = True

    row_count = max(len(original_lines), len(revised_lines), len(reasons))
    table = doc.add_table(rows=row_count + 1, cols=3)
    table.style = "Light Grid Accent 1"

    headers = ["原文", "修改后", "修改原因"]
    for i, h_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h_text)
        _set_chinese_font(run, size=10)
        run.bold = True

    for i in range(row_count):
        row = table.rows[i + 1]
        for j, data_list in enumerate(
            [original_lines, revised_lines, reasons]
        ):
            cell = row.cells[j]
            cell.text = ""
            text = data_list[i] if i < len(data_list) else ""
            run = cell.paragraphs[0].add_run(text)
            _set_chinese_font(run, size=9)

    # Tips
    if tips:
        doc.add_paragraph("")
        h = doc.add_heading(level=2)
        run = h.add_run("提升建议")
        _set_chinese_font(run, size=14)
        run.bold = True
        for i, tip in enumerate(tips, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {tip}")
            _set_chinese_font(run, size=11)
            run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)

    path = os.path.join(OUTPUT_DIR, filename)
    doc.save(path)
    return path
